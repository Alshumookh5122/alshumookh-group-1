from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "ALSHUMOOKH_Security_Compliance_Certification_Roadmap.docx"

INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5F6B7A"
LIGHT_BLUE = "E8F1FA"
LIGHT_GRAY = "F2F4F7"
GREEN = "1F7A4D"
GOLD = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = "D8DEE9", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            margins = tc_pr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side, value in (("top", "100"), ("bottom", "100"), ("start", "140"), ("end", "140")):
                node = margins.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    margins.append(node)
                node.set(qn("w:w"), value)
                node.set(qn("w:type"), "dxa")


def set_run(run, *, size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold


def add_para(doc: Document, text: str = "", *, style: str | None = None, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run(run, bold=True, color=INK)
        run = p.add_run(text[len(bold_prefix):])
        set_run(run, color=INK)
    else:
        run = p.add_run(text)
        set_run(run, color=INK)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6 if level <= 2 else 4)
    for run in p.runs:
        set_run(run, size=16 if level == 1 else 13 if level == 2 else 12, color=BLUE if level <= 2 else DARK_BLUE, bold=True)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE, accent: str = DARK_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    set_cell_borders(cell, "C8D7E6")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run(r, size=11, color=accent, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(body)
    set_run(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        set_cell_fill(hdr[idx], LIGHT_GRAY)
        set_cell_borders(hdr[idx])
        p = hdr[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        set_run(r, size=10, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_borders(cells[idx])
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(text)
            color = INK
            if text in {"Immediate", "Required", "High", "In progress"}:
                color = RED if text in {"Required", "High"} else GOLD
            if text in {"Ready", "Low", "Implemented"}:
                color = GREEN
            set_run(r, size=9.5, color=color, bold=text in {"Required", "High", "Ready", "Implemented", "Immediate"})
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.1)
    r = p.add_run(text)
    set_run(r, size=10.5, color=INK)


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]
    header.text = "ALSHUMOOKH GLOBAL | Security & Compliance Roadmap"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_run(run, size=9, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.text = "Confidential operational planning document"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_run(run, size=9, color=MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("ALSHUMOOKH Security, Compliance & Certification Roadmap")
    set_run(r, size=24, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("Production domain: https://api.alshumookh-pay.com | Prepared: " + date.today().isoformat())
    set_run(r, size=11.5, color=MUTED, bold=True)

    meta = [
        ["Document Type", "Certification roadmap and counterparty assurance pack"],
        ["Platform Scope", "API-to-API settlement receiver, MoonPay payment links, Alchemy on-chain verification, audit logging"],
        ["Current Status", "Production API reachable; certification program to be initiated"],
        ["Audience", "Internal management, external sender/counterparty, auditors, compliance consultants"],
    ]
    add_table(doc, ["Field", "Value"], meta, [2500, 6860])

    add_callout(
        doc,
        "Executive position",
        "The platform can operate while formal certifications are being processed, provided that high-value flows use controlled counterparties, verified wallets, signed payloads, Alchemy verification, and manual approval for exceptional cases. The certification roadmap below converts the current production system into an externally auditable financial technology control environment.",
    )

    add_heading(doc, "1. Immediate Assurance Package", 1)
    add_para(doc, "These are the items that can be provided to the sender immediately while formal third-party certifications are being obtained.")
    add_table(
        doc,
        ["Item", "Purpose", "Status", "Evidence"],
        [
            ["Production API endpoint", "Confirms the receiving system is live and reachable.", "Ready", "https://api.alshumookh-pay.com/api/v1/payloads/schema"],
            ["TLS/HTTPS domain", "Encrypts traffic in transit and validates the production hostname.", "Ready", "HTTPS active on api.alshumookh-pay.com"],
            ["HMAC/OAuth/JWS/JWE-ready API", "Supports authenticated, signed, and encrypted sender payloads.", "Implemented", "Settlement ingest endpoint and schema"],
            ["Audit log trail", "Records API requests, auth events, security events, and payment operations.", "Implemented", "Admin dashboard and database audit_logs"],
            ["Alchemy verification", "Validates transaction hash, receiver wallet, chain status, asset, and amount.", "Implemented", "Ethereum/Base RPC verification pipeline"],
            ["MoonPay integration", "Creates payment links through configured MoonPay Commerce credentials.", "Implemented", "MoonPay provider service and redirect flow"],
        ],
        [2100, 3100, 1300, 2860],
    )

    add_heading(doc, "2. Certification Priorities", 1)
    add_table(
        doc,
        ["Certification / Control", "Why it matters", "Priority", "Estimated sequence"],
        [
            ["Penetration Test Report", "Fastest independent proof that the public API, auth, WAF, and dashboard have been tested.", "Immediate", "Start within 1 week"],
            ["PCI-DSS scope memo / SAQ", "Clarifies that card data is handled by MoonPay and is not processed or stored by ALSHUMOOKH.", "Immediate", "Prepare before large card-funded flow"],
            ["AML/KYC Policy Pack", "Defines onboarding, sanctions checks, transaction monitoring, and escalation rules.", "Required", "Start immediately"],
            ["SOC 2 Type I", "Independent report over security controls at a point in time.", "High", "30-60 days after evidence collection"],
            ["ISO 27001 readiness", "Information security management system, policies, risk register, and control ownership.", "High", "60-120 days depending on auditor"],
            ["SOC 2 Type II", "Longer audit proving controls operated over time.", "High", "3-12 month observation period"],
            ["VASP / crypto regulatory review", "Determines licensing/registration obligations for digital asset activity.", "Required", "Legal counsel review before scaling"],
            ["Data Protection / Privacy Pack", "Defines privacy notice, retention, access control, and data subject handling.", "Required", "Parallel with AML/KYC"],
        ],
        [2200, 3650, 1300, 2210],
    )

    add_heading(doc, "3. Operational Control Matrix", 1)
    add_table(
        doc,
        ["Control domain", "Current platform control", "Certification evidence to prepare"],
        [
            ["Access control", "Admin dashboard protected by admin session/API key; client API keys and OAuth credentials supported.", "Access policy, admin list, key rotation records, least-privilege review"],
            ["Network security", "Cloudflare-compatible WAF behavior, suspicious path blocking, rate limits, scanner detection, security headers.", "WAF rule summary, event exports, incident response procedure"],
            ["Payload security", "HMAC signatures, OAuth bearer tokens, JWS/JWE-ready payload handling, idempotency keys.", "Integration guide, sample signed payload, key management procedure"],
            ["Settlement verification", "Alchemy RPC verification for Ethereum/Base, receiver wallet match, amount/asset checks, manual review states.", "Verification screenshots, sample tx hash, reconciliation SOP"],
            ["Auditability", "API request logs, security events, admin operations, payment orders, Alchemy events.", "Audit export procedure, retention policy, evidence samples"],
            ["Payment provider boundary", "MoonPay creates hosted payment links; sensitive card data should not transit through ALSHUMOOKH servers.", "Provider agreement, PCI scope memo, data-flow diagram"],
            ["Business continuity", "Render deployment, health checks, fallback documentation, production domain monitoring.", "Backup/restore SOP, uptime monitoring, incident contact list"],
        ],
        [2300, 3550, 3510],
    )

    add_heading(doc, "4. 30/60/90 Day Execution Plan", 1)
    add_table(
        doc,
        ["Phase", "Target outcome", "Actions"],
        [
            ["Day 0-7", "Counterparty assurance and safe pilot", "Deliver API guide, schema, signed payload samples, wallet confirmation, MoonPay/Alchemy screenshots, and run one low-value test."],
            ["Day 8-30", "External security validation", "Commission penetration test, close findings, document WAF/rate-limit/security headers, complete PCI scope memo."],
            ["Day 31-60", "Policy and evidence base", "Build AML/KYC policy, privacy notice, access control policy, change management log, vendor register, risk register."],
            ["Day 61-90", "Audit readiness", "Begin SOC 2 Type I or ISO 27001 readiness assessment, collect control evidence, prepare management assertion."],
            ["90+ days", "Institutional-grade certification path", "Proceed with SOC 2 Type II observation, ISO 27001 certification audit, and VASP/legal registration path if required."],
        ],
        [1400, 2500, 5460],
    )

    add_heading(doc, "5. Required Environment Evidence", 1)
    add_para(doc, "The following production values should be confirmed through Render screenshots or an environment evidence export. Secrets must be redacted before sharing externally.")
    for item in [
        "PUBLIC_BASE_URL = https://api.alshumookh-pay.com",
        "MOONPAY_API_KEY, MOONPAY_API_SECRET, MOONPAY_DEPOSIT_ID, MOONPAY_WEBHOOK_SECRET present and redacted",
        "ALCHEMY_WEBHOOK_SIGNING_KEY and Ethereum/Base RPC configuration present and redacted",
        "MASTER_WALLET_ETHEREUM and MASTER_WALLET_BASE present, validated, and approved by management",
        "ADMIN_API_KEY and client API keys present, rotated, and not using placeholder values",
        "DATABASE_URL configured for production PostgreSQL, with backups enabled at provider level",
        "CORS_ALLOWED_ORIGINS restricted to approved domains",
        "HEALTH_ALLOWED_IPS or HEALTHCHECK_TOKEN configured for internal health checks",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6. Sender-Facing Technical Proof", 1)
    add_table(
        doc,
        ["Proof", "How to provide it", "Disclosure level"],
        [
            ["Live API schema", "Share GET https://api.alshumookh-pay.com/api/v1/payloads/schema", "Public"],
            ["Signed request sample", "Provide HMAC/OAuth sample with dummy secrets", "Controlled"],
            ["Low-value test transaction", "Sender submits payload with tx_hash; ALSHUMOOKH returns review/verification status", "Controlled"],
            ["Wallet ownership proof", "Provide management-signed wallet address confirmation and optional explorer link", "Controlled"],
            ["MoonPay proof", "Show successful creation of a small payment link without exposing credentials", "Controlled"],
            ["Alchemy proof", "Show transaction verification result and audit event without exposing API key", "Controlled"],
        ],
        [2450, 4110, 2800],
    )

    add_heading(doc, "7. Key Risks and Mitigations", 1)
    add_table(
        doc,
        ["Risk", "Impact", "Mitigation"],
        [
            ["Formal certifications are not yet issued", "Some institutional senders may not accept high-value transfers.", "Provide assurance pack now; start pen test and SOC/ISO readiness immediately."],
            ["Regulatory classification depends on jurisdiction", "Licensing may be required for VASP/payment activity.", "Engage legal counsel and document operating model before scaling."],
            ["MoonPay/card-data scope misunderstood", "Counterparty may request PCI-DSS unnecessarily.", "Document that card data is hosted by MoonPay; prepare PCI scope memo."],
            ["Wrong network or wallet", "Funds could be delayed or unrecoverable.", "Use approved master wallets, confirmation screen, and Alchemy receiver match checks."],
            ["High-value transfer without pilot", "Operational and reconciliation risk.", "Require low-value end-to-end test before large settlement."],
        ],
        [2400, 2750, 4210],
    )

    add_heading(doc, "8. Recommended Next Step", 1)
    add_callout(
        doc,
        "Decision request",
        "Proceed with the assurance package immediately, then book a third-party penetration test as the first external certificate-style proof. In parallel, prepare AML/KYC, PCI scope, and SOC 2/ISO 27001 readiness documentation.",
        fill="FFF7E6",
        accent=GOLD,
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Appendix A: Certification Roadmap Checklist", 1)
    checklist_rows = [
        ["☐", "Confirm all production environment variables in Render"],
        ["☐", "Export redacted environment evidence"],
        ["☐", "Run low-value sender test payload"],
        ["☐", "Run low-value MoonPay payment link test"],
        ["☐", "Run Alchemy tx_hash verification evidence capture"],
        ["☐", "Prepare PCI scope memo"],
        ["☐", "Prepare AML/KYC and sanctions screening policy"],
        ["☐", "Book penetration test"],
        ["☐", "Create SOC 2 / ISO 27001 evidence folder"],
        ["☐", "Begin VASP/legal regulatory classification review"],
    ]
    add_table(doc, ["Done", "Action"], checklist_rows, [900, 8460])

    OUT.unlink(missing_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
