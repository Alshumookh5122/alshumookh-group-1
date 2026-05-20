from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_FILES = [
    ROOT / "ENTERPRISE_UPGRADE_MASTER_PLAN.md",
    ROOT / "RELEASE_CHECKLIST_ENTERPRISE.md",
    ROOT / "ZERO_BREAK_DEPLOYMENT_PLAN.md",
]
OUTPUT_FILE = ROOT / "ALSHUMOOKH_Enterprise_Upgrade_Plan.docx"


BRAND_DARK = RGBColor(15, 23, 42)
BRAND_GOLD = RGBColor(191, 148, 59)
TEXT_MUTED = RGBColor(71, 85, 105)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def parse_markdown(path: Path):
    sections = []
    current = None

    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("# "):
            current = {"title": line[2:].strip(), "items": []}
            sections.append(current)
            continue

        if line.startswith("## "):
            if current is None:
                current = {"title": path.stem, "items": []}
                sections.append(current)
            current["items"].append(("h2", line[3:].strip()))
            continue

        if line.startswith("### "):
            if current is None:
                current = {"title": path.stem, "items": []}
                sections.append(current)
            current["items"].append(("h3", line[4:].strip()))
            continue

        if line.startswith("- [ ] "):
            current["items"].append(("check", line[6:].strip()))
            continue

        if line.startswith("- "):
            current["items"].append(("bullet", line[2:].strip()))
            continue

        numbered = False
        if len(line) > 3 and line[0].isdigit():
            for i, ch in enumerate(line):
                if ch == "." and i > 0 and line[:i].isdigit():
                    current["items"].append(("number", line[i + 1 :].strip()))
                    numbered = True
                    break
                if not ch.isdigit():
                    break
        if numbered:
            continue

        current["items"].append(("para", line))

    return sections


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def build_cover(doc: Document):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ALSHUMOOKH\nEnterprise Upgrade Plan")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = BRAND_DARK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Master plan, release checklist, and zero-break deployment guide"
    )
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_MUTED

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Date", "2026-05-11"),
        ("Primary Domain", "https://api.alshumookh-pay.com"),
        ("Prepared For", "ALSHUMOOKH internal enterprise upgrade cycle"),
    ]
    for (label, value), row in zip(rows, table.rows):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], "0F172A")
        for run in row.cells[0].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    doc.add_paragraph("")

    note = doc.add_paragraph()
    note.style = "Intense Quote"
    note.add_run(
        "This document consolidates the planned enterprise upgrade path for the "
        "current production platform while preserving existing sender, admin, and client flows."
    )

    doc.add_page_break()


def build_summary(doc: Document):
    heading = doc.add_paragraph()
    heading.style = "Heading 1"
    heading.add_run("Executive Summary")

    for text in [
        "The current platform is already operational and production-facing. "
        "This plan upgrades the existing system into a stronger enterprise-style posture "
        "without disrupting active production workflows.",
        "The document is organized into three operational parts: the master upgrade plan, "
        "the release checklist, and the zero-break deployment plan.",
    ]:
        p = doc.add_paragraph(text)
        p.style = "Body Text"

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Part", "Purpose", "Outcome"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "BF943B")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    rows = [
        ("Master Plan", "Define scope, goals, phases, and constraints", "Shared execution baseline"),
        ("Release Checklist", "Verify readiness before and after release", "Lower deployment risk"),
        ("Zero-Break Plan", "Control production rollout and rollback path", "Operational continuity"),
    ]
    for part, purpose, outcome in rows:
        cells = table.add_row().cells
        cells[0].text = part
        cells[1].text = purpose
        cells[2].text = outcome

    doc.add_page_break()


def apply_styles(doc: Document):
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.color.rgb = BRAND_DARK

    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(12.5)
    styles["Heading 3"].font.size = Pt(11)


def add_section_content(doc: Document, title: str, items):
    h1 = doc.add_paragraph()
    h1.style = "Heading 1"
    h1.add_run(title)

    for kind, text in items:
        if kind == "h2":
            p = doc.add_paragraph()
            p.style = "Heading 2"
            p.add_run(text)
        elif kind == "h3":
            p = doc.add_paragraph()
            p.style = "Heading 3"
            p.add_run(text)
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(text)
        elif kind == "check":
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("□ " + text)
        elif kind == "number":
            p = doc.add_paragraph(style="List Number")
            p.add_run(text)
        else:
            p = doc.add_paragraph(text)
            p.style = "Body Text"


def build_doc():
    doc = Document()
    apply_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    build_cover(doc)
    build_summary(doc)

    for path in SOURCE_FILES:
        for section_data in parse_markdown(path):
            add_section_content(doc, section_data["title"], section_data["items"])
            doc.add_paragraph("")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("ALSHUMOOKH Enterprise Upgrade Plan - Page ")
    add_page_number(footer)

    doc.save(OUTPUT_FILE)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT_FILE)
