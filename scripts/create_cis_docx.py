from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "ALSHUMOOKH_CIS_API_TO_API.docx"
LOGO = ROOT / "scripts" / "alshumookh_logo_mark.png"


BLUE = "1F5FD0"
DARK = "111827"
MUTED = "667085"
LINE = "D8E0EA"
LIGHT = "F4F7FB"
GOLD = "B9892F"


def make_logo() -> None:
    img = Image.new("RGB", (420, 420), "#111827")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((18, 18, 402, 402), radius=34, fill="#111827", outline="#B9892F", width=10)
    try:
        font_big = ImageFont.truetype("Arial Bold.ttf", 132)
        font_small = ImageFont.truetype("Arial.ttf", 34)
    except Exception:
        font_big = ImageFont.load_default()
        font_small = ImageFont.load_default()
    draw.text((210, 168), "AS", anchor="mm", font=font_big, fill="#FFFFFF")
    draw.text((210, 286), "ALSHUMOOKH", anchor="mm", font=font_small, fill="#D9B56B")
    img.save(LOGO)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = LINE) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text: str, bold: bool = False, color: str = DARK, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, rows: list[tuple[str, str]], widths: tuple[float, float] = (2.25, 4.25)) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(widths[0])
        cells[1].width = Inches(widths[1])
        set_cell_text(cells[0], label, bold=True, color=MUTED)
        set_cell_text(cells[1], value, color=DARK)
        set_cell_shading(cells[0], LIGHT)
        set_cell_border(cells[0])
        set_cell_border(cells[1])


def add_header_row(table, labels: list[str]) -> None:
    row = table.rows[0]
    for i, label in enumerate(labels):
        set_cell_text(row.cells[i], label, bold=True, color="FFFFFF", size=9)
        set_cell_shading(row.cells[i], BLUE)
        set_cell_border(row.cells[i], BLUE)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string(BLUE if level == 1 else DARK)


def add_note(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF8E8")
    set_cell_border(cell, "E6C06A")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(DARK)


def build_doc() -> None:
    make_logo()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12)

    header = section.header.paragraphs[0]
    header.text = "ALSHUMOOKH GLOBAL BANKING FINANCE & CREDIT | Client Information Sheet"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Confidential - For authorized integration use only"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.name = "Arial"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    cover = doc.add_table(rows=1, cols=2)
    cover.alignment = WD_TABLE_ALIGNMENT.LEFT
    cover.autofit = False
    cover.cell(0, 0).width = Inches(1.35)
    cover.cell(0, 1).width = Inches(5.7)
    cover.cell(0, 0).paragraphs[0].add_run().add_picture(str(LOGO), width=Inches(1.12))
    p = cover.cell(0, 1).paragraphs[0]
    title_run = p.add_run("CLIENT INFORMATION SHEET (CIS)")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string(DARK)
    p2 = cover.cell(0, 1).add_paragraph()
    r = p2.add_run("ALSHUMOOKH GLOBAL BANKING FINANCE & CREDIT")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p3 = cover.cell(0, 1).add_paragraph()
    r = p3.add_run("Payment Receiving & API-to-API Integration Information")
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    for cell in cover.row_cells(0):
        set_cell_border(cell, "FFFFFF")

    doc.add_paragraph()
    add_table(
        doc,
        [
            ("Environment", "Production"),
            ("Service Type", "Fiat-to-Crypto & Crypto Receiving Gateway"),
            ("Primary Provider", "Coinbase Onramp"),
            ("Blockchain Monitoring", "Alchemy"),
            ("Settlement Wallet Type", "Ledger Wallet"),
            ("Production API", "https://api.alshumookh-pay.com"),
        ],
    )

    add_heading(doc, "1. Company Information")
    add_table(
        doc,
        [
            ("Company Name", "ALSHUMOOKH GLOBAL BANKING FINANCE & CREDIT"),
            ("Service Description", "API-to-API payment receiving gateway for fiat-to-crypto transactions and direct crypto settlement monitoring."),
            ("Receiving Model", "Client API Request -> ALSHUMOOKH API -> Coinbase Checkout -> Coinbase Confirmation -> Ledger Wallet Settlement"),
        ],
    )

    add_heading(doc, "2. Production API Details")
    add_table(
        doc,
        [
            ("Base URL", "https://api.alshumookh-pay.com"),
            ("Create Transaction", "POST https://api.alshumookh-pay.com/api/v1/transactions"),
            ("Check Status", "GET https://api.alshumookh-pay.com/api/v1/transactions/{transaction_id}"),
            ("System Health", "GET https://api.alshumookh-pay.com/health"),
        ],
    )

    add_heading(doc, "3. Client Authentication")
    add_table(
        doc,
        [
            ("Header", "X-API-Key: CLIENT_API_KEY_HERE"),
            ("Idempotency", "Idempotency-Key: UNIQUE_PAYMENT_REFERENCE"),
            ("Content Type", "Content-Type: application/json"),
            ("Security Rule", "Client API Key must be stored securely on the sender backend/server and must not be exposed publicly."),
        ],
    )

    add_heading(doc, "4. Create Payment Request")
    add_table(
        doc,
        [
            ("external_id", "PAYMENT-1001"),
            ("network", "base"),
            ("fiat_currency", "USD"),
            ("crypto_currency", "USDC"),
            ("fiat_amount", "100"),
            ("country", "US"),
            ("subdivision", "CA"),
            ("redirect_url", "https://sender-domain.com/payment-success"),
        ],
    )
    add_note(doc, "The API returns a Coinbase checkout_url. The sender must open or redirect the payer to that URL to complete payment.")

    add_heading(doc, "5. Supported Payment Details")
    add_table(
        doc,
        [
            ("Supported Fiat Currency", "USD"),
            ("Supported Crypto Asset", "USDC"),
            ("Primary Network", "Base"),
            ("Alternative Network", "Ethereum"),
            ("Payment Methods", "Debit/Credit Card, Bank Transfer, Coinbase Balance, USDC Wallet, and other Coinbase-supported methods where available."),
        ],
    )

    add_heading(doc, "6. ALSHUMOOKH Ledger Wallet")
    add_table(
        doc,
        [
            ("Wallet Address", "0xBD682cfD8382a90adfDd6745780D3D7959c4d939"),
            ("Supported Use", "Coinbase Onramp settlement, direct USDC crypto transfer, and Alchemy blockchain monitoring."),
            ("Direct Transfer", "Asset: USDC | Network: Base or Ethereum | Destination: ALSHUMOOKH Ledger Wallet"),
        ],
    )
    add_note(doc, "Send only supported assets on the correct network. Wrong-network transfers may be unrecoverable.")

    add_heading(doc, "7. Fiat / Bank Transfer Through Coinbase")
    add_table(
        doc,
        [
            ("Recipient Name", "Coinbase Singapore Pte. Ltd."),
            ("Recipient Address", "One Marina Boulevard, 28-00, 018989, Singapore"),
            ("Bank Name", "STANDARD CHARTERED BANK (SINGAPORE) LIMITED"),
            ("SWIFT Code", "SCBLSG22XXX"),
            ("Bank Account Number", "99200334405"),
            ("Bank Address", "MARINA BOULEVARD, 8, MARINA BAY FINANCIAL CENTRE, 27, 01"),
            ("Reference", "NA / As provided by Coinbase checkout"),
        ],
    )
    add_note(doc, "Bank transfer details must be followed only when displayed by Coinbase for the payer's checkout session. Preferred flow: Create Transaction API -> Coinbase checkout_url.")

    add_heading(doc, "8. Alchemy Monitoring")
    add_table(
        doc,
        [
            ("Monitored Wallet", "0xBD682cfD8382a90adfDd6745780D3D7959c4d939"),
            ("Monitored Networks", "Base, Ethereum"),
            ("Monitored Asset", "USDC"),
            ("Usage", "Wallet activity monitoring, on-chain payment detection, transaction confirmation tracking, webhook confirmation."),
        ],
    )

    add_heading(doc, "9. Transaction Status Values")
    status_table = doc.add_table(rows=1, cols=2)
    status_table.style = "Table Grid"
    add_header_row(status_table, ["Status", "Description"])
    for status, description in [
        ("PENDING", "Payment created and awaiting payer action."),
        ("PROCESSING", "Payment or blockchain confirmation in progress."),
        ("COMPLETED", "Payment confirmed successfully."),
        ("FAILED", "Payment failed."),
        ("EXPIRED", "Payment expired."),
        ("REFUNDED", "Payment refunded."),
    ]:
        cells = status_table.add_row().cells
        set_cell_text(cells[0], status, bold=True)
        set_cell_text(cells[1], description)
        set_cell_border(cells[0])
        set_cell_border(cells[1])

    add_heading(doc, "10. Security Restrictions")
    add_table(
        doc,
        [
            ("Provided to Sender", "Client API Key, API Base URL, Create Transaction Endpoint, Status Endpoint, Ledger Wallet Address."),
            ("Not Provided", "Admin API Key, Coinbase API Secret, Coinbase Webhook Secret, Alchemy Secret, Database Credentials, Private Wallet Keys, Ledger Private Keys."),
        ],
    )

    add_heading(doc, "11. Production Test")
    add_table(
        doc,
        [
            ("external_id", "TEST-PAYMENT-001"),
            ("Idempotency-Key", "TEST-PAYMENT-001"),
            ("network", "base"),
            ("fiat_currency", "USD"),
            ("crypto_currency", "USDC"),
            ("fiat_amount", "10"),
        ],
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("End of Client Information Sheet")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(MUTED)

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
